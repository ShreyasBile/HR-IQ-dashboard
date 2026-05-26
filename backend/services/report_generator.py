import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.analytics import get_analytics
from services.anomaly_engine import get_anomalies
from services.benchmarks import get_benchmarks

def generate_report(sections: list):
    doc = Document()
    
    # Title
    title = doc.add_heading('HRIQ Intelligence Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Analytics Data
    analytics = get_analytics()
    if "error" not in analytics and "Executive Summary" in sections:
        doc.add_heading('Executive Summary', level=1)
        insights = analytics.get('insights', [])
        for insight in insights:
            doc.add_paragraph(insight, style='List Bullet')
            
        # Overview Stats
        doc.add_heading('Workforce Overview', level=2)
        overview = analytics.get('overview', {})
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        for key, val in overview.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace('avg', 'Average ').replace('resignationRate', 'Resignation Rate (%)')
            row_cells[1].text = str(val)

    # Anomalies Data
    if "Anomalies" in sections:
        anomalies = get_anomalies()
        if "error" not in anomalies:
            doc.add_heading('Detected Anomalies', level=1)
            for alert in anomalies.get('data', []):
                p = doc.add_paragraph()
                run = p.add_run(f"[{alert['severity']}] {alert['title']}")
                run.bold = True
                doc.add_paragraph(alert['description'])
                p = doc.add_paragraph()
                p.add_run(f"Recommended Action: {alert['action']}").italic = True

    # Benchmarking Data
    if "Benchmarking" in sections:
        benchmarks = get_benchmarks("Technology") # Defaulting to Technology for report
        if "error" not in benchmarks:
            doc.add_heading(f'Industry Benchmarking (Technology)', level=1)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Company Value'
            hdr_cells[2].text = 'Industry Median'
            hdr_cells[3].text = 'Status'
            
            for item in benchmarks.get('data', []):
                row_cells = table.add_row().cells
                row_cells[0].text = item['metric']
                row_cells[1].text = str(item['company'])
                row_cells[2].text = str(item['median'])
                row_cells[3].text = item['status']

    # Save Document
    os.makedirs("data/reports", exist_ok=True)
    report_path = "data/reports/HRIQ_Report.docx"
    doc.save(report_path)
    
    return report_path
